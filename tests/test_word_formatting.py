from __future__ import annotations

import json
import base64
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
PNG_1X1 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="


def run_format_docx(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, "-X", "utf8", str(ROOT / "skills" / "word-formatting" / "scripts" / "format_docx.py"), *args],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=True,
    )


def run_inspect_docx(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, "-X", "utf8", str(ROOT / "skills" / "word-formatting" / "scripts" / "inspect_docx.py"), *args],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=True,
    )


def load_template(path: Path) -> dict:
    run_format_docx("--write-template", str(path))
    return json.loads(path.read_text(encoding="utf-8"))


def visible_paragraph_text(paragraph) -> str:
    return "".join(paragraph._p.xpath(".//w:t/text()"))


def mark_repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "1")
    trpr.append(header)


class WordFormattingTests(unittest.TestCase):
    def test_default_template_uses_required_page_margins_and_cleanup_features(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            config_path = Path(tmp) / "config.json"
            config = load_template(config_path)
            self.assertEqual(config["page"]["top_cm"], 2.5)
            self.assertEqual(config["page"]["bottom_cm"], 2.5)
            self.assertEqual(config["page"]["left_cm"], 3.0)
            self.assertEqual(config["page"]["right_cm"], 2.5)
            self.assertTrue(config["features"]["protect_formulas"])
            self.assertTrue(config["features"]["normalize_cjk_latin_spacing"])
            self.assertTrue(config["features"]["remove_unused_styles"])
            self.assertEqual(config["formula"]["output_format"], "MathML")
            self.assertEqual(config["formula"]["parameter_output_format"], "MathML")
            self.assertEqual(config["formula"]["body_parameter_output_format"], "MathML")
            self.assertEqual(config["figures"]["table_image_width_cm"], 7.7)
            self.assertTrue(config["figures"]["preserve_aspect_ratio"])
            self.assertTrue(config["captions"]["use_word_caption_fields"])
            self.assertEqual(config["captions"]["figure_seq_id"], "Figure")
            self.assertEqual(config["captions"]["table_seq_id"], "Table")
            self.assertTrue(config["styles"]["h1"]["q_format"])
            self.assertEqual(config["styles"]["h1"]["outline_level"], 0)
            self.assertEqual(config["styles"]["h2"]["outline_level"], 1)
            self.assertEqual(config["styles"]["h3"]["outline_level"], 2)
            self.assertEqual(config["styles"]["h1"]["line_spacing"], 1.0)
            self.assertEqual(config["styles"]["h2"]["line_spacing"], 1.0)
            self.assertEqual(config["styles"]["h3"]["line_spacing"], 1.0)
            self.assertEqual(config["styles"]["h1"]["ui_priority"], 10)
            self.assertTrue(config["styles"]["h1"]["keep_next"])
            self.assertTrue(config["styles"]["h1"]["keep_lines"])

    def test_format_removes_cross_run_technical_term_and_number_unit_spaces(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "mixed.docx"
            config_path = Path(tmp) / "config.json"
            config = load_template(config_path)
            config_path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")

            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("\u672c\u6587\u7814\u7a76")
            paragraph.add_run(" UHPC ")
            paragraph.add_run("\u6e7f\u63a5\u7f1d\u4e0e")
            paragraph.add_run(" RC ")
            paragraph.add_run("\u57fa\u51c6\uff0c\u5c3a\u5bf8\u4e3a")
            paragraph.add_run("1000 mm")
            paragraph.add_run("\uff0c\u53e6\u6709")
            paragraph.add_run(" 400mm ")
            paragraph.add_run("\u8bd5\u4ef6\uff0c\u5f3a\u5ea6\u4e3a")
            paragraph.add_run("40 MPa")
            paragraph.add_run("\uff0c\u5899\u9ad8")
            paragraph.add_run("51.6 m ")
            paragraph.add_run("\u9ad8\uff0c\u5c3a\u5bf8\u4e3a")
            paragraph.add_run("1500 \u00d7 400 \u00d7 3000 mm")
            paragraph.add_run("\uff0c\u6bcf\u6392 8 \u6839\uff0c\u56fe 1 \u6240\u793a\uff0c\u8868 6 \u7d2f\u8ba1\u8017\u80fd")
            paragraph.add_run("\u3002")
            doc.save(docx)

            result = run_format_docx("--docx", str(docx), "--config", str(config_path), "--json")
            report = json.loads(result.stdout)
            text = Document(docx).paragraphs[0].text

            self.assertIn("UHPC\u6e7f\u63a5\u7f1d", text)
            self.assertIn("RC\u57fa\u51c6", text)
            self.assertIn("1000mm", text)
            self.assertIn("400mm\u8bd5\u4ef6", text)
            self.assertIn("40MPa", text)
            self.assertIn("51.6m\u9ad8", text)
            self.assertIn("1500\u00d7400\u00d73000mm", text)
            self.assertIn("\u6bcf\u63928\u6839", text)
            self.assertIn("\u56fe1\u6240\u793a", text)
            self.assertIn("\u88686\u7d2f\u8ba1\u8017\u80fd", text)
            self.assertNotIn(" UHPC ", text)
            self.assertNotIn(" RC ", text)
            self.assertNotIn(" 400mm ", text)
            self.assertNotIn("1000 mm", text)
            self.assertNotIn("40 MPa", text)
            self.assertNotIn("1500 \u00d7 400 \u00d7 3000 mm", text)
            self.assertGreaterEqual(report["cjk_alnum_spaces_removed"], 12)
            self.assertTrue(report["formula_hash_ok"])
            self.assertEqual(report["formula_output_format"], "MathML")
            self.assertEqual(report["formula_parameter_output_format"], "MathML")
            self.assertEqual(report["body_parameter_output_format"], "MathML")

    def test_format_removes_unused_custom_styles_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "styles.docx"
            config_path = Path(tmp) / "config.json"
            config = load_template(config_path)
            config_path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")

            doc = Document()
            doc.styles.add_style("Unused Custom Style", WD_STYLE_TYPE.PARAGRAPH)
            doc.add_paragraph("\u6b63\u6587 UHPC \u5185\u5bb9\u3002")
            self.assertIn("Unused Custom Style", {style.name for style in doc.styles})
            doc.save(docx)

            result = run_format_docx("--docx", str(docx), "--config", str(config_path), "--json")
            report = json.loads(result.stdout)
            styles = {style.name for style in Document(docx).styles}

            self.assertNotIn("Unused Custom Style", styles)
            self.assertIn("Unused Custom Style", report["unused_styles_removed"])

    def test_inspect_reports_mathml_formula_and_body_parameter_output_formats(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "inspect.docx"
            config_path = Path(tmp) / "config.json"
            config = load_template(config_path)
            config_path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")

            doc = Document()
            doc.add_paragraph("\u6b63\u6587\u53c2\u6570 x \u548c y \u9700\u6309 MathML \u8f93\u51fa\u3002")
            doc.save(docx)

            result = run_inspect_docx("--docx", str(docx), "--config", str(config_path), "--json")
            report = json.loads(result.stdout)

            self.assertEqual(report["formula_output_format"], "MathML")
            self.assertEqual(report["formula_parameter_output_format"], "MathML")
            self.assertEqual(report["body_parameter_output_format"], "MathML")

    def test_format_writes_office_style_metadata_for_headings(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "metadata.docx"
            config_path = Path(tmp) / "config.json"
            config = load_template(config_path)
            config_path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")

            doc = Document()
            doc.add_paragraph("Title")
            doc.add_paragraph("1 Introduction")
            doc.add_paragraph("1.1 Background")
            doc.add_paragraph("1.1.1 Scope")
            doc.save(docx)

            result = run_format_docx("--docx", str(docx), "--config", str(config_path), "--json")
            report = json.loads(result.stdout)
            formatted = Document(docx)
            h1 = formatted.styles[config["styles"]["h1"]["name"]]
            h2 = formatted.styles[config["styles"]["h2"]["name"]]
            h3 = formatted.styles[config["styles"]["h3"]["name"]]

            self.assertEqual(h1.element.xpath("./w:qFormat")[0].get(qn("w:val")), "1")
            self.assertEqual(h1.element.xpath("./w:uiPriority")[0].get(qn("w:val")), "10")
            self.assertEqual(h1.element.xpath("./w:pPr/w:outlineLvl")[0].get(qn("w:val")), "0")
            self.assertEqual(h1.element.xpath("./w:pPr/w:keepNext")[0].get(qn("w:val")), "1")
            self.assertEqual(h1.element.xpath("./w:pPr/w:keepLines")[0].get(qn("w:val")), "1")
            self.assertEqual(h2.element.xpath("./w:pPr/w:outlineLvl")[0].get(qn("w:val")), "1")
            self.assertEqual(h3.element.xpath("./w:pPr/w:outlineLvl")[0].get(qn("w:val")), "2")
            self.assertEqual(report["style_office_metadata_issues"], [])

    def test_format_sets_figure_table_image_width_to_7_7_cm(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "figure-width.docx"
            image = Path(tmp) / "image.png"
            config_path = Path(tmp) / "config.json"
            config = load_template(config_path)
            config_path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")
            image.write_bytes(base64.b64decode(PNG_1X1))

            doc = Document()
            table = doc.add_table(rows=2, cols=1)
            table.cell(0, 0).paragraphs[0].add_run().add_picture(str(image), width=Cm(3.0))
            table.cell(1, 0).text = "Figure 1 Test image"
            doc.save(docx)

            result = run_format_docx("--docx", str(docx), "--config", str(config_path), "--json")
            report = json.loads(result.stdout)
            formatted = Document(docx)
            inline = formatted.tables[0].cell(0, 0)._tc.xpath(".//wp:inline")[0]
            width_cm = round(int(inline.xpath("./wp:extent")[0].get("cx")) / int(Cm(1)), 2)

            self.assertEqual(width_cm, 7.7)
            self.assertEqual(report["figure_table_image_widths_cm"], [7.7])

    def test_format_removes_repeat_header_only_from_figure_tables_and_centers_regular_table_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "table-row-options.docx"
            image = Path(tmp) / "image.png"
            config_path = Path(tmp) / "config.json"
            config = load_template(config_path)
            config_path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")
            image.write_bytes(base64.b64decode(PNG_1X1))

            doc = Document()
            figure_table = doc.add_table(rows=2, cols=1)
            figure_table.cell(0, 0).paragraphs[0].add_run().add_picture(str(image), width=Cm(3.0))
            figure_table.cell(1, 0).text = "Figure 1 Test image"
            mark_repeat_header(figure_table.rows[0])

            data_table = doc.add_table(rows=2, cols=2)
            data_table.cell(0, 0).text = "A"
            data_table.cell(0, 1).text = "B"
            data_table.cell(1, 0).text = "1"
            data_table.cell(1, 1).text = "2"
            mark_repeat_header(data_table.rows[0])
            doc.save(docx)

            result = run_format_docx("--docx", str(docx), "--config", str(config_path), "--json")
            report = json.loads(result.stdout)
            formatted = Document(docx)
            formatted_figure = formatted.tables[0]
            formatted_data = formatted.tables[1]

            self.assertFalse(formatted_figure.rows[0]._tr.xpath("./w:trPr/w:tblHeader"))
            self.assertTrue(formatted_data.rows[0]._tr.xpath("./w:trPr/w:tblHeader"))
            for row in formatted_data.rows:
                for cell in row.cells:
                    self.assertEqual(cell._tc.xpath("./w:tcPr/w:vAlign/@w:val"), ["center"])
            self.assertEqual(report["figure_table_repeating_header_row_count"], 0)
            self.assertEqual(report["regular_table_vertical_alignment_issues"], [])

    def test_format_converts_figure_and_table_captions_to_seq_fields(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "caption-fields.docx"
            config_path = Path(tmp) / "config.json"
            config = load_template(config_path)
            config_path.write_text(json.dumps(config, ensure_ascii=False), encoding="utf-8")

            doc = Document()
            doc.add_paragraph("Figure 3 Test image")
            doc.add_paragraph("\u8868 6 \u7d2f\u8ba1\u8017\u80fd")
            doc.save(docx)

            result = run_format_docx("--docx", str(docx), "--config", str(config_path), "--json")
            report = json.loads(result.stdout)
            formatted = Document(docx)
            text = [visible_paragraph_text(paragraph) for paragraph in formatted.paragraphs]
            instr_values = formatted._body._element.xpath(".//w:fldSimple/@w:instr")

            self.assertIn("Figure 3 Test image", text)
            self.assertIn("\u88686\u7d2f\u8ba1\u8017\u80fd", text)
            self.assertTrue(any("SEQ Figure" in value for value in instr_values))
            self.assertTrue(any("SEQ Table" in value for value in instr_values))
            self.assertEqual(report["caption_fields_converted"], 2)
            self.assertEqual(report["caption_seq_field_counts"], {"Figure": 1, "Table": 1})


if __name__ == "__main__":
    unittest.main()
