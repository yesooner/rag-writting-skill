from __future__ import annotations

import argparse
import json
import re
import shutil
from copy import deepcopy
from datetime import datetime
from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


DEFAULT_CITATION_PATTERN = r"\[\d+(?:\s*[-\u2013,\uff0c]\s*\d+)*\]"
CJK_ALNUM_SPACING_PATTERN = re.compile(r"(?<=[\u3400-\u9fff])\s+(?=[A-Za-z0-9])|(?<=[A-Za-z0-9])\s+(?=[\u3400-\u9fff])")
CJK_LATIN_SPACING_PATTERN = CJK_ALNUM_SPACING_PATTERN
FONT_ALIASES = {
    "宋体": "SimSun",
    "黑体": "SimHei",
    "仿宋": "FangSong",
    "楷体": "KaiTi",
}
ALIGNMENT = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def default_config() -> dict:
    return {
        "page": {
            "width_cm": 21.0,
            "height_cm": 29.7,
            "top_cm": 2.5,
            "bottom_cm": 2.5,
            "left_cm": 3.0,
            "right_cm": 2.5,
            "header_cm": 1.5,
            "footer_cm": 1.5,
        },
        "header_footer": {
            "enabled": True,
            "header_text": "document_stem",
            "page_number": True,
        },
        "features": {
            "wrap_figures": True,
            "format_tables": True,
            "superscript_citations": True,
            "protect_formulas": True,
            "normalize_cjk_latin_spacing": True,
        },
        "citation_pattern": DEFAULT_CITATION_PATTERN,
        "styles": {
            "title": style("Title", "SimHei", "Times New Roman", 18, True, "center", 1.5, 0, 0, 0),
            "subtitle": style("Subtitle", "SimSun", "Times New Roman", 14, True, "center", 1.15, 0, 18, 0),
            "abstract_head": style("Abstract Heading", "SimHei", "Times New Roman", 10.5, True, "left", 1.0, 0, 0, 0),
            "abstract_body": style("Abstract Body", "SimSun", "Times New Roman", 10.5, False, "justify", 1.5, 0, 0, 21),
            "keywords": style("Keywords", "SimSun", "Times New Roman", 10.5, False, "left", 1.0, 0, 0, 0),
            "h1": style("Heading 1 Custom", "SimHei", "Times New Roman", 14, True, "left", 1.5, 0, 0, 0),
            "h2": style("Heading 2 Custom", "SimSun", "Times New Roman", 12, False, "left", 1.5, 0, 0, 0),
            "h3": style("Heading 3 Custom", "SimSun", "Times New Roman", 12, False, "left", 1.5, 0, 0, 0),
            "body": style("Body Text Custom", "SimSun", "Times New Roman", 12, False, "justify", 1.5, 0, 0, 24),
            "figure": style("Figure", "SimSun", "Times New Roman", 12, False, "center", 1.0, 0, 0, 0),
            "figure_caption": style("Figure Caption", "SimSun", "Times New Roman", 10.5, True, "center", 1.0, 0, 0, 0),
            "table_caption": style("Table Caption", "SimSun", "Times New Roman", 10.5, True, "center", 1.0, 0, 0, 0),
            "table_text": style("Table Text", "SimSun", "Times New Roman", 10.5, False, "center", 1.0, 0, 0, 0),
            "reference": style("Reference Text", "SimSun", "Times New Roman", 10.5, False, "justify", 1.0, 0, 0, 0, hanging_indent_pt=21),
            "formula": style("Formula", "SimSun", "Times New Roman", 12, False, "center", 1.0, 3, 3, 0),
        },
        "patterns": {
            "figure_caption": [r"^图\s*\d+", r"^Figure\s+\d+\s+(?!shows\b)"],
            "table_caption": [r"^表\s*\d+", r"^Table\s+\d+"],
            "reference": [r"^\[\d+\]"],
            "h3": [r"^\d+\.\d+\.\d+\s+"],
            "h2": [r"^\d+\.\d+\s+"],
            "h1": [r"^\d+\s+"],
            "abstract_head": [r"^摘要$", r"^Abstract$"],
            "keywords": [r"^关键词", r"^Keywords"],
        },
    }


def style(
    name: str,
    font_cn: str,
    font_en: str,
    size_pt: float,
    bold: bool,
    alignment: str,
    line_spacing: float,
    space_before_pt: float,
    space_after_pt: float,
    first_line_indent_pt: float,
    hanging_indent_pt: float | None = None,
) -> dict:
    return {
        "name": name,
        "font_cn": font_cn,
        "font_en": font_en,
        "size_pt": size_pt,
        "bold": bold,
        "alignment": alignment,
        "line_spacing": line_spacing,
        "space_before_pt": space_before_pt,
        "space_after_pt": space_after_pt,
        "first_line_indent_pt": first_line_indent_pt,
        "hanging_indent_pt": hanging_indent_pt,
    }


def load_config(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def normalize_font(name: str) -> str:
    return FONT_ALIASES.get(name, name)


def math_hashes(doc: Document) -> list[str]:
    return [sha256(etree.tostring(node)).hexdigest() for node in doc._body._element.xpath(".//m:oMath | .//m:oMathPara")]


def compiled_patterns(config: dict, key: str) -> list[re.Pattern]:
    return [re.compile(pattern) for pattern in config.get("patterns", {}).get(key, [])]


def text_matches(patterns: list[re.Pattern], text: str) -> bool:
    return any(pattern.search(text) for pattern in patterns)


def alignment(value: str):
    return ALIGNMENT.get(str(value).lower(), WD_ALIGN_PARAGRAPH.LEFT)


def ensure_rfonts_on_rpr(rpr, font_cn: str, font_en: str) -> None:
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr, value in (("w:ascii", font_en), ("w:hAnsi", font_en), ("w:eastAsia", font_cn), ("w:cs", font_en)):
        rfonts.set(qn(attr), value)


def ensure_rfonts(run, font_cn: str, font_en: str) -> None:
    run.font.name = font_en
    ensure_rfonts_on_rpr(run._element.get_or_add_rPr(), font_cn, font_en)


def create_styles(doc: Document, config: dict) -> dict[str, object]:
    styles = {}
    for role, spec in config.get("styles", {}).items():
        name = spec["name"]
        if name in doc.styles:
            doc_style = doc.styles[name]
        else:
            doc_style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            if "Normal" in doc.styles:
                doc_style.base_style = doc.styles["Normal"]
        font_cn = normalize_font(spec.get("font_cn", "SimSun"))
        font_en = normalize_font(spec.get("font_en", "Times New Roman"))
        doc_style.font.name = font_en
        doc_style.font.size = Pt(spec.get("size_pt", 12))
        doc_style.font.bold = bool(spec.get("bold", False))
        ensure_rfonts_on_rpr(doc_style.element.get_or_add_rPr(), font_cn, font_en)
        apply_paragraph_format(doc_style.paragraph_format, spec)
        styles[role] = doc_style
    return styles


def apply_paragraph_format(paragraph_format, spec: dict) -> None:
    paragraph_format.alignment = alignment(spec.get("alignment", "left"))
    paragraph_format.space_before = Pt(spec.get("space_before_pt", 0))
    paragraph_format.space_after = Pt(spec.get("space_after_pt", 0))
    paragraph_format.line_spacing = spec.get("line_spacing", 1.0)
    hanging = spec.get("hanging_indent_pt")
    if hanging is not None:
        paragraph_format.left_indent = Pt(abs(hanging))
        paragraph_format.first_line_indent = Pt(-abs(hanging))
    else:
        paragraph_format.left_indent = Pt(0)
        paragraph_format.first_line_indent = Pt(spec.get("first_line_indent_pt", 0))


def apply_style(paragraph, styles: dict[str, object], config: dict, role: str) -> None:
    if role not in styles:
        return
    spec = config["styles"][role]
    paragraph.style = styles[role]
    apply_paragraph_format(paragraph.paragraph_format, spec)
    paragraph.alignment = alignment(spec.get("alignment", "left"))
    font_cn = normalize_font(spec.get("font_cn", "SimSun"))
    font_en = normalize_font(spec.get("font_en", "Times New Roman"))
    for run in paragraph.runs:
        ensure_rfonts(run, font_cn, font_en)
        run.font.size = Pt(spec.get("size_pt", 12))
        run.bold = bool(spec.get("bold", False))


def set_page_setup(doc: Document, config: dict) -> None:
    page = config.get("page", {})
    for section in doc.sections:
        if "width_cm" in page:
            section.page_width = Cm(page["width_cm"])
        if "height_cm" in page:
            section.page_height = Cm(page["height_cm"])
        if "top_cm" in page:
            section.top_margin = Cm(page["top_cm"])
        if "bottom_cm" in page:
            section.bottom_margin = Cm(page["bottom_cm"])
        if "left_cm" in page:
            section.left_margin = Cm(page["left_cm"])
        if "right_cm" in page:
            section.right_margin = Cm(page["right_cm"])
        if "header_cm" in page:
            section.header_distance = Cm(page["header_cm"])
        if "footer_cm" in page:
            section.footer_distance = Cm(page["footer_cm"])


def clear_paragraph(paragraph) -> None:
    element = paragraph._p
    for child in list(element):
        element.remove(child)


def set_header_footer(doc: Document, config: dict, path: Path) -> None:
    hf = config.get("header_footer", {})
    if not hf.get("enabled", False):
        return
    header_text = hf.get("header_text", "document_stem")
    if header_text == "document_stem":
        header_text = path.stem
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(header_text)
        ensure_rfonts(run, "SimSun", "Times New Roman")
        run.font.size = Pt(9)
        add_bottom_border(paragraph)
        if hf.get("page_number", True):
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            clear_paragraph(fp)
            add_page_number(fp)


def add_bottom_border(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(" ")
    ensure_rfonts(run, "SimSun", "Times New Roman")
    run.font.size = Pt(10)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)
    paragraph.add_run(" ")


def classify_paragraph(paragraph, config: dict, non_empty_index: int) -> str:
    text = paragraph.text.strip()
    patterns = config.get("patterns", {})
    for role in ("figure_caption", "table_caption", "reference", "h3", "h2", "h1", "abstract_head", "keywords"):
        if text_matches([re.compile(p) for p in patterns.get(role, [])], text):
            return role
    if paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//w:pict"):
        return "figure"
    if paragraph._p.xpath(".//m:oMath") or paragraph._p.xpath(".//m:oMathPara"):
        return "formula" if not text else "body"
    if non_empty_index == 1 and "title" in config.get("styles", {}):
        return "title"
    return "body"


def style_top_level_paragraphs(doc: Document, styles: dict[str, object], config: dict) -> dict[str, int]:
    counts: dict[str, int] = {}
    non_empty_index = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            non_empty_index += 1
        role = classify_paragraph(paragraph, config, non_empty_index)
        apply_style(paragraph, styles, config, role)
        counts[role] = counts.get(role, 0) + 1
    return counts


def set_table_borders(table, visible: bool) -> None:
    tblpr = table._tbl.tblPr
    if tblpr is None:
        tblpr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblpr)
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    edges = (
        [("top", "single", "12"), ("bottom", "single", "12"), ("insideH", "single", "6"), ("left", "nil", "0"), ("right", "nil", "0"), ("insideV", "nil", "0")]
        if visible
        else [(edge, "nil", "0") for edge in ("top", "bottom", "insideH", "left", "right", "insideV")]
    )
    for edge, value, size in edges:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_margins(cell, value: str = "80") -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for name in ("top", "start", "bottom", "end"):
        node = tcmar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcmar.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


def is_figure_caption(text: str, config: dict) -> bool:
    return text_matches(compiled_patterns(config, "figure_caption"), text.strip())


def compiled_patterns(config: dict, role: str) -> list[re.Pattern]:
    return [re.compile(pattern) for pattern in config.get("patterns", {}).get(role, [])]


def wrap_figures(doc: Document, styles: dict[str, object], config: dict) -> int:
    body = doc._body._element
    pairs = []
    for paragraph_element in list(body.findall(qn("w:p"))):
        if not (paragraph_element.xpath(".//w:drawing") or paragraph_element.xpath(".//w:pict")):
            continue
        next_element = paragraph_element.getnext()
        if next_element is not None and next_element.tag == qn("w:p"):
            caption_text = "".join(next_element.xpath(".//w:t/text()")).strip()
            if is_figure_caption(caption_text, config):
                pairs.append((paragraph_element, next_element))
    wrapped = 0
    for image_p, caption_p in pairs:
        table = doc.add_table(rows=2, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, visible=False)
        body.remove(table._tbl)
        body.insert(body.index(image_p), table._tbl)
        table.cell(0, 0)._tc.clear_content()
        table.cell(1, 0)._tc.clear_content()
        table.cell(0, 0)._tc.append(deepcopy(image_p))
        table.cell(1, 0)._tc.append(deepcopy(caption_p))
        body.remove(image_p)
        body.remove(caption_p)
        wrapped += 1
    format_tables(doc, styles, config)
    return wrapped


def format_tables(doc: Document, styles: dict[str, object], config: dict) -> tuple[int, int]:
    figure_count = 0
    regular_count = 0
    for table in doc.tables:
        is_figure_table = (
            len(table.rows) == 2
            and len(table.columns) == 1
            and bool(table.cell(0, 0)._tc.xpath(".//w:drawing") or table.cell(0, 0)._tc.xpath(".//w:pict"))
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if is_figure_table:
            set_table_borders(table, visible=False)
            for cell in table._cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, "0")
            for paragraph in table.cell(0, 0).paragraphs:
                apply_style(paragraph, styles, config, "figure")
            for paragraph in table.cell(1, 0).paragraphs:
                apply_style(paragraph, styles, config, "figure_caption")
            figure_count += 1
        else:
            set_table_borders(table, visible=True)
            for row_index, row in enumerate(table.rows):
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell)
                    for paragraph in cell.paragraphs:
                        apply_style(paragraph, styles, config, "table_text")
                        if row_index == 0:
                            for run in paragraph.runs:
                                run.bold = True
            regular_count += 1
    return figure_count, regular_count


def iter_all_paragraphs(doc: Document):
    seen = set()
    for paragraph in doc.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    key = id(paragraph._p)
                    if key not in seen:
                        seen.add(key)
                        yield paragraph


def normalize_cjk_alnum_spacing_text(text: str) -> str:
    return CJK_ALNUM_SPACING_PATTERN.sub("", text)


def normalize_cjk_latin_spacing_text(text: str) -> str:
    return normalize_cjk_alnum_spacing_text(text)


def cjk_alnum_spacing_issue_count(text: str) -> int:
    return len(CJK_ALNUM_SPACING_PATTERN.findall(text or ""))


def cjk_latin_spacing_issue_count(text: str) -> int:
    return cjk_alnum_spacing_issue_count(text)


def normalize_cjk_latin_spacing(doc: Document) -> int:
    changed = 0
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            text = run.text or ""
            normalized = normalize_cjk_latin_spacing_text(text)
            if normalized != text:
                changed += cjk_alnum_spacing_issue_count(text)
                run.text = normalized
    return changed


def document_text(doc: Document) -> str:
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def normalized_document_text(doc: Document) -> str:
    return "\n".join(normalize_cjk_latin_spacing_text(paragraph.text) for paragraph in doc.paragraphs)


def make_text_run(text: str, source_run, superscript: bool):
    new_run = OxmlElement("w:r")
    old_rpr = source_run.find(qn("w:rPr"))
    if old_rpr is not None:
        new_run.append(deepcopy(old_rpr))
    rpr = new_run.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        new_run.insert(0, rpr)
    for va in list(rpr.findall(qn("w:vertAlign"))):
        rpr.remove(va)
    if superscript:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)
    node = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace() or "  " in text:
        node.set(qn("xml:space"), "preserve")
    node.text = text
    new_run.append(node)
    return new_run


def is_reference_paragraph(paragraph, config: dict) -> bool:
    ref_style = config.get("styles", {}).get("reference", {}).get("name")
    return (ref_style and paragraph.style.name == ref_style) or text_matches(compiled_patterns(config, "reference"), paragraph.text.strip())


def superscript_citations(doc: Document, config: dict) -> int:
    citation_re = re.compile(config.get("citation_pattern", DEFAULT_CITATION_PATTERN))
    changed = 0
    for paragraph in list(iter_all_paragraphs(doc)):
        if is_reference_paragraph(paragraph, config) or not citation_re.search(paragraph.text):
            continue
        for run in list(paragraph.runs):
            text = run.text or ""
            if not citation_re.search(text):
                continue
            if run._r.xpath(".//w:drawing") or run._r.xpath(".//w:pict") or run._r.xpath(".//m:oMath") or run._r.xpath(".//m:oMathPara"):
                continue
            parts = []
            last = 0
            for match in citation_re.finditer(text):
                if match.start() > last:
                    parts.append((text[last : match.start()], False))
                parts.append((match.group(0), True))
                last = match.end()
                changed += 1
            if last < len(text):
                parts.append((text[last:], False))
            parent = run._r.getparent()
            position = parent.index(run._r)
            for offset, (part, superscript) in enumerate([p for p in parts if p[0]]):
                parent.insert(position + offset, make_text_run(part, run._r, superscript))
            parent.remove(run._r)
    return changed


def collect_report(doc: Document, target: Path, backup: Path | None, config: dict, formula_hash_ok: bool, extra: dict) -> dict:
    styles = [spec["name"] for spec in config.get("styles", {}).values() if spec.get("name")]
    existing_styles = {style.name for style in doc.styles}
    citation_re = re.compile(config.get("citation_pattern", DEFAULT_CITATION_PATTERN))
    body_total = body_super = ref_total = ref_super = 0
    for paragraph in iter_all_paragraphs(doc):
        is_ref = is_reference_paragraph(paragraph, config)
        for run in paragraph.runs:
            found = citation_re.findall(run.text or "")
            if not found:
                continue
            is_super = bool(run._r.xpath('./w:rPr/w:vertAlign[@w:val="superscript"]'))
            if is_ref:
                ref_total += len(found)
                if is_super:
                    ref_super += len(found)
            else:
                body_total += len(found)
                if is_super:
                    body_super += len(found)
    figure_tables = sum(
        1
        for table in doc.tables
        if len(table.rows) == 2
        and len(table.columns) == 1
        and bool(table.cell(0, 0)._tc.xpath(".//w:drawing") or table.cell(0, 0)._tc.xpath(".//w:pict"))
    )
    return {
        "target": str(target),
        "backup": str(backup) if backup else None,
        "missing_configured_styles": [name for name in styles if name not in existing_styles],
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "image_count": len(doc.inline_shapes),
        "figure_table_count": figure_tables,
        "top_level_image_paragraph_count": sum(1 for p in doc.paragraphs if p._p.xpath(".//w:drawing") or p._p.xpath(".//w:pict")),
        "body_citation_count": body_total,
        "body_citation_superscript_count": body_super,
        "reference_number_count": ref_total,
        "reference_number_superscript_count": ref_super,
        "formula_node_count": len(doc._body._element.xpath(".//m:oMath | .//m:oMathPara")),
        "formula_hash_ok": formula_hash_ok,
        **extra,
    }


def format_docx(path: Path, config: dict) -> dict:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = path.with_name(path.stem + f"_backup_before_word_format_{stamp}" + path.suffix)
    shutil.copy2(path, backup)
    doc = Document(str(path))
    initial_math = math_hashes(doc)
    initial_images = len(doc.inline_shapes)
    normalize_spacing_enabled = config.get("features", {}).get("normalize_cjk_latin_spacing", False)
    initial_text = normalized_document_text(doc) if normalize_spacing_enabled else document_text(doc)

    styles = create_styles(doc, config)
    set_page_setup(doc, config)
    set_header_footer(doc, config, path)
    style_counts = style_top_level_paragraphs(doc, styles, config)
    wrapped = wrap_figures(doc, styles, config) if config.get("features", {}).get("wrap_figures", False) else 0
    figure_tables, regular_tables = format_tables(doc, styles, config) if config.get("features", {}).get("format_tables", False) else (0, len(doc.tables))
    cjk_latin_spaces_removed = normalize_cjk_latin_spacing(doc) if normalize_spacing_enabled else 0
    citation_marks = superscript_citations(doc, config) if config.get("features", {}).get("superscript_citations", False) else 0

    if config.get("features", {}).get("protect_formulas", True) and math_hashes(doc) != initial_math:
        raise RuntimeError("Formula XML changed; aborting without saving.")
    if len(doc.inline_shapes) != initial_images:
        raise RuntimeError(f"Image count changed {initial_images} -> {len(doc.inline_shapes)}; aborting without saving.")
    final_text = document_text(doc)
    if final_text != initial_text:
        raise RuntimeError("Top-level paragraph text changed; aborting without saving.")
    doc.save(str(path))
    saved = Document(str(path))
    return collect_report(
        saved,
        path,
        backup,
        config,
        formula_hash_ok=True,
        extra={
            "style_application_counts": style_counts,
            "wrapped_figures": wrapped,
            "formatted_figure_tables": figure_tables,
            "formatted_regular_tables": regular_tables,
            "citation_marks_changed": citation_marks,
            "cjk_alnum_spaces_removed": cjk_latin_spaces_removed,
            "cjk_latin_spaces_removed": cjk_latin_spaces_removed,
        },
    )


def main() -> int:
    parser = argparse.ArgumentParser(description="Format a .docx using a confirmed JSON style config.")
    parser.add_argument("--docx", type=Path)
    parser.add_argument("--config", type=Path)
    parser.add_argument("--write-template", type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()

    if args.write_template:
        args.write_template.write_text(json.dumps(default_config(), ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"template_written: {args.write_template}")
        return 0
    if not args.docx or not args.config:
        raise SystemExit("--docx and --config are required unless --write-template is used")

    report = format_docx(args.docx, load_config(args.config))
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        for key, value in report.items():
            print(f"{key}: {value}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
