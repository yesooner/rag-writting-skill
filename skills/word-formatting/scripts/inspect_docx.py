from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn


DEFAULT_CITATION_PATTERN = r"\[\d+(?:\s*[-\u2013,\uff0c]\s*\d+)*\]"
CJK_ALNUM_SPACING_PATTERN = re.compile(r"(?<=[\u3400-\u9fff])\s+(?=[A-Za-z0-9])|(?<=[A-Za-z0-9])\s+(?=[\u3400-\u9fff])")
CJK_LATIN_SPACING_PATTERN = CJK_ALNUM_SPACING_PATTERN
EMU_PER_CM = int(Cm(1))


def load_config(path: Path | None) -> dict:
    if path is None:
        return default_config()
    return json.loads(path.read_text(encoding="utf-8"))


def default_config() -> dict:
    return {
        "styles": {
            "figure_caption": {"name": "Figure Caption"},
            "reference": {"name": "Reference Text"},
        },
        "patterns": {
            "reference": [r"^\[\d+\]"],
        },
        "formula": {
            "output_format": "MathML",
            "parameter_output_format": "MathML",
            "body_parameter_output_format": "MathML",
        },
        "captions": {
            "use_word_caption_fields": True,
            "figure_seq_id": "Figure",
            "table_seq_id": "Table",
            "skip_existing_seq_fields": True,
            "preserve_caption_text": True,
        },
    }


def compiled_patterns(config: dict, key: str) -> list[re.Pattern]:
    return [re.compile(pattern) for pattern in config.get("patterns", {}).get(key, [])]


def text_matches(patterns: list[re.Pattern], text: str) -> bool:
    return any(pattern.search(text) for pattern in patterns)


def iter_all_paragraphs(doc: Document):
    seen = set()
    for paragraph in doc.paragraphs:
        key = id(paragraph._p)
        if key not in seen:
            seen.add(key)
            yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    key = id(paragraph._p)
                    if key not in seen:
                        seen.add(key)
                        yield paragraph


def is_superscript_run(run) -> bool:
    return bool(run._r.xpath('./w:rPr/w:vertAlign[@w:val="superscript"]'))


def cjk_alnum_spacing_issue_count(text: str) -> int:
    return len(CJK_ALNUM_SPACING_PATTERN.findall(text or ""))


def cjk_latin_spacing_issue_count(text: str) -> int:
    return cjk_alnum_spacing_issue_count(text)


def formula_output_format(config: dict) -> str:
    return str(config.get("formula", {}).get("output_format", "MathML"))


def formula_parameter_output_format(config: dict) -> str:
    return str(config.get("formula", {}).get("parameter_output_format", "MathML"))


def body_parameter_output_format(config: dict) -> str:
    return str(config.get("formula", {}).get("body_parameter_output_format", "MathML"))


def child_val(element, path: str) -> str | None:
    found = element.xpath(path)
    if not found:
        return None
    return found[0].get(qn("w:val"))


def expected_on_off(value: bool) -> str:
    return "1" if value else "0"


def style_office_metadata_issues(doc: Document, config: dict) -> list[dict]:
    issues = []
    existing_styles = {style.name: style for style in doc.styles}
    required_typography = ["size_pt", "line_spacing", "alignment"]
    for role, spec in config.get("styles", {}).items():
        name = spec.get("name")
        if not name:
            continue
        missing_fields = [field for field in required_typography if spec.get(field) is None]
        if missing_fields:
            issues.append({"role": role, "style": name, "missing_config_fields": missing_fields})
        doc_style = existing_styles.get(name)
        if doc_style is None:
            continue
        element = doc_style.element
        if "q_format" in spec and child_val(element, "./w:qFormat") != expected_on_off(bool(spec.get("q_format"))):
            issues.append({"role": role, "style": name, "missing_or_mismatched": "qFormat"})
        if spec.get("ui_priority") is not None and child_val(element, "./w:uiPriority") != str(int(spec["ui_priority"])):
            issues.append({"role": role, "style": name, "missing_or_mismatched": "uiPriority"})
        if spec.get("outline_level") is not None and child_val(element, "./w:pPr/w:outlineLvl") != str(int(spec["outline_level"])):
            issues.append({"role": role, "style": name, "missing_or_mismatched": "outlineLvl"})
        if "keep_next" in spec and child_val(element, "./w:pPr/w:keepNext") != expected_on_off(bool(spec.get("keep_next"))):
            issues.append({"role": role, "style": name, "missing_or_mismatched": "keepNext"})
        if "keep_lines" in spec and child_val(element, "./w:pPr/w:keepLines") != expected_on_off(bool(spec.get("keep_lines"))):
            issues.append({"role": role, "style": name, "missing_or_mismatched": "keepLines"})
    return issues


def figure_table_image_widths_cm(doc: Document) -> list[float]:
    widths = []
    for table in doc.tables:
        is_figure_table = (
            len(table.rows) == 2
            and len(table.columns) == 1
            and bool(table.cell(0, 0)._tc.xpath(".//w:drawing") or table.cell(0, 0)._tc.xpath(".//w:pict"))
        )
        if not is_figure_table:
            continue
        for inline in table.cell(0, 0)._tc.xpath(".//wp:inline"):
            extents = inline.xpath("./wp:extent")
            if not extents:
                continue
            width = int(extents[0].get("cx", "0"))
            widths.append(round(width / EMU_PER_CM, 2))
    return widths


def caption_seq_field_counts(doc: Document) -> dict[str, int]:
    counts = {"Figure": 0, "Table": 0}
    for value in doc._body._element.xpath(".//w:fldSimple/@w:instr | .//w:instrText/text()"):
        text = str(value)
        if "SEQ Figure" in text:
            counts["Figure"] += 1
        if "SEQ Table" in text:
            counts["Table"] += 1
    return counts


def inspect_docx(path: Path, config: dict) -> dict:
    doc = Document(str(path))
    citation_re = re.compile(config.get("citation_pattern", DEFAULT_CITATION_PATTERN))
    ref_patterns = compiled_patterns(config, "reference")
    ref_style = config.get("styles", {}).get("reference", {}).get("name")
    required_styles = [spec["name"] for spec in config.get("styles", {}).values() if spec.get("name")]
    existing_styles = {style.name for style in doc.styles}

    body_total = 0
    body_super = 0
    ref_total = 0
    ref_super = 0
    body_unsup = []
    mixed_spacing_total = 0
    mixed_spacing_examples = []
    for index, paragraph in enumerate(iter_all_paragraphs(doc), 1):
        text = paragraph.text.strip()
        spacing_count = cjk_alnum_spacing_issue_count(text)
        if spacing_count:
            mixed_spacing_total += spacing_count
            if len(mixed_spacing_examples) < 10:
                mixed_spacing_examples.append(
                    {
                        "paragraph_index": index,
                        "style": paragraph.style.name,
                        "issue_count": spacing_count,
                        "text": text[:160],
                    }
                )
        is_ref = (ref_style and paragraph.style.name == ref_style) or text_matches(ref_patterns, text)
        for run in paragraph.runs:
            found = citation_re.findall(run.text or "")
            if not found:
                continue
            if is_ref:
                ref_total += len(found)
                if is_superscript_run(run):
                    ref_super += len(found)
            else:
                body_total += len(found)
                if is_superscript_run(run):
                    body_super += len(found)
                else:
                    body_unsup.append(
                        {
                            "paragraph_index": index,
                            "style": paragraph.style.name,
                            "marks": found,
                            "text": text[:160],
                        }
                    )

    figure_tables = 0
    regular_tables = 0
    caption_sizes = []
    for table in doc.tables:
        is_figure_table = (
            len(table.rows) == 2
            and len(table.columns) == 1
            and bool(table.cell(0, 0)._tc.xpath(".//w:drawing") or table.cell(0, 0)._tc.xpath(".//w:pict"))
        )
        if is_figure_table:
            figure_tables += 1
            for paragraph in table.cell(1, 0).paragraphs:
                for run in paragraph.runs:
                    if run.font.size:
                        caption_sizes.append(run.font.size.pt)
        else:
            regular_tables += 1

    return {
        "target": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "regular_table_count": regular_tables,
        "image_count": len(doc.inline_shapes),
        "formula_node_count": len(doc._body._element.xpath(".//m:oMath | .//m:oMathPara")),
        "formula_output_format": formula_output_format(config),
        "formula_parameter_output_format": formula_parameter_output_format(config),
        "body_parameter_output_format": body_parameter_output_format(config),
        "style_office_metadata_issues": style_office_metadata_issues(doc, config),
        "missing_configured_styles": [name for name in required_styles if name not in existing_styles],
        "figure_table_count": figure_tables,
        "figure_table_image_widths_cm": figure_table_image_widths_cm(doc),
        "caption_seq_field_counts": caption_seq_field_counts(doc),
        "top_level_image_paragraph_count": sum(
            1 for paragraph in doc.paragraphs if paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//w:pict")
        ),
        "figure_caption_size_unique": sorted(set(caption_sizes)),
        "body_citation_count": body_total,
        "body_citation_superscript_count": body_super,
        "body_citation_unsuperscript_count": len(body_unsup),
        "body_citation_unsuperscript_examples": body_unsup[:10],
        "reference_number_count": ref_total,
        "reference_number_superscript_count": ref_super,
        "cjk_alnum_spacing_issue_count": mixed_spacing_total,
        "cjk_alnum_spacing_examples": mixed_spacing_examples,
        "cjk_latin_spacing_issue_count": mixed_spacing_total,
        "cjk_latin_spacing_examples": mixed_spacing_examples,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Inspect .docx formatting using an optional JSON style config.")
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--config", type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()

    report = inspect_docx(args.docx, load_config(args.config))
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        for key, value in report.items():
            print(f"{key}: {value}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
